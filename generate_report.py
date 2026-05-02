"""
generate_report.py
Generates project_report.docx for:
  Smart HR Operations & Workforce Management System
  Group 16 — SE2002 SDA, FAST NUCES Spring 2026

Run:
    pip install python-docx
    python generate_report.py

Diagrams are written as Mermaid code inside yellow-highlighted boxes.
Give the output .docx to Claude.ai and ask:
  "Replace every DIAGRAM PLACEHOLDER block with the rendered Mermaid diagram as an image."
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    return p

def add_diagram_placeholder(doc, title, diagram_type, mermaid_code):
    """
    Adds a yellow-highlighted box with Mermaid code.
    Claude.ai will replace this with a rendered diagram image.
    """
    doc.add_paragraph()
    # Border table trick for a "box"
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'FFF9C4')  # light yellow

    p1 = cell.add_paragraph()
    r1 = p1.add_run('[DIAGRAM PLACEHOLDER - FOR CLAUDE.AI]')
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
    r1.font.size = Pt(11)

    p2 = cell.add_paragraph()
    r2 = p2.add_run(f'Title: {title}')
    r2.bold = True
    r2.font.size = Pt(10)

    p3 = cell.add_paragraph()
    r3 = p3.add_run(f'Type: {diagram_type}')
    r3.font.size = Pt(10)

    p4 = cell.add_paragraph()
    r4 = p4.add_run('Mermaid Code (render this as a diagram image):')
    r4.italic = True
    r4.font.size = Pt(9)

    p5 = cell.add_paragraph()
    r5 = p5.add_run(mermaid_code)
    r5.font.name = 'Courier New'
    r5.font.size = Pt(8)

    p6 = cell.add_paragraph()
    r6 = p6.add_run('>> Claude.ai instruction: Render the Mermaid code above as a UML diagram image and insert it here. Remove this text block.')
    r6.italic = True
    r6.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
    r6.font.size = Pt(9)

    doc.add_paragraph()

def add_simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_bg(hdr_cells[i], '1F4E79')
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for run in row_cells[i].paragraphs[0].runs:
                run.font.size = Pt(10)

    return table

def page_break(doc):
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
# MERMAID DIAGRAM STRINGS
# ─────────────────────────────────────────────────────────────

USE_CASE_DIAGRAM = """
graph TD
    HR[HR Officer]
    Admin[Admin]
    Emp[Employee]

    HR  --> UC01[UC-01 Register New Employee]
    HR  --> UC02[UC-02 Update Employee Record]
    HR  --> UC03[UC-03 Assign to Department]
    HR  --> UC04[UC-04 Initiate Offboarding]
    HR  --> UC05[UC-05 View Employee Profile]
    HR  --> UC07[UC-07 Approve or Reject Leave]
    HR  --> UC09[UC-09 Approve Attendance Correction]
    HR  --> UC12[UC-12 Submit Performance Evaluation]
    HR  --> UC13[UC-13 Monitor Probation Status]

    Emp --> UC05
    Emp --> UC06[UC-06 Submit Leave Request]
    Emp --> UC08[UC-08 Record Daily Attendance]
    Emp --> UC09
    Emp --> UC10[UC-10 View Leave Balance & Dashboard]

    Admin --> UC03
    Admin --> UC11[UC-11 Initiate Evaluation Cycle]
    Admin --> UC14[UC-14 Generate Compliance Report]
    Admin --> UC15[UC-15 Generate HR Analytics Report]
"""

DOMAIN_MODEL = """
classDiagram
    class Employee {
        +int id
        +String firstName
        +String lastName
        +String email
        +String nationalId
        +String employmentType
        +BigDecimal basicSalary
        +String status
        +getFullName() String
    }
    class Department {
        +int id
        +String name
        +int managerId
        +int maxHeadcount
        +int currentHeadcount
    }
    class Designation {
        +int id
        +String title
        +String salaryGrade
    }
    class LeaveRequest {
        +int id
        +String leaveType
        +LocalDate startDate
        +LocalDate endDate
        +int daysRequested
        +String status
    }
    class LeaveBalance {
        +int annualBalance
        +int sickBalance
        +int personalBalance
    }
    class AttendanceRecord {
        +LocalDate attendanceDate
        +LocalDateTime checkInTime
        +LocalDateTime checkOutTime
        +double totalHours
        +String attendanceStatus
    }
    class Payroll {
        +BigDecimal basicSalary
        +BigDecimal taxDeduction
        +BigDecimal benefitsDeduction
        +BigDecimal netPay
    }
    class OffboardingWorkflow {
        +String separationType
        +LocalDate lastWorkingDate
        +String status
        +validateNoticePeriod() boolean
    }
    class ProbationRecord {
        +LocalDate startDate
        +LocalDate endDate
        +String decision
        +String status
    }
    class PerformanceEvaluation {
        +double aggregateScore
        +String remarks
        +String status
    }
    class ComplianceReport {
        +String reportType
        +String format
        +String archivePath
        +getReportType() String
        +getSummary() String
    }
    class HRAnalyticsReport {
        +String reportPeriod
        +String summaryMetrics
        +getReportType() String
        +getSummary() String
    }
    class Report {
        <<abstract>>
        +int id
        +LocalDateTime generatedAt
        +getReportType()* String
        +getSummary()* String
        +getFormat()* String
    }

    Employee "many" --> "1" Department
    Employee "many" --> "1" Designation
    Employee "1" --> "1" LeaveBalance
    Employee "1" --> "many" LeaveRequest
    Employee "1" --> "many" AttendanceRecord
    Employee "1" --> "many" Payroll
    Employee "1" --> "0..1" OffboardingWorkflow
    Employee "1" --> "0..1" ProbationRecord
    Employee "1" --> "many" PerformanceEvaluation
    Report <|-- ComplianceReport
    Report <|-- HRAnalyticsReport
"""

SSD_UC01 = """
sequenceDiagram
    actor HRO as HR Officer
    participant SYS as :System

    HRO ->> SYS: navigateToRegisterModule()
    SYS -->> HRO: displayRegistrationForm()

    HRO ->> SYS: enterPersonalDetails(name, nationalID, dob, contact)
    SYS -->> HRO: duplicateCheckResult()

    HRO ->> SYS: selectDepartmentDesignation(dept, designation, empType, joiningDate)
    SYS -->> HRO: validationResult() + probationPeriod()

    HRO ->> SYS: uploadDocuments(contractFile, idScan)
    SYS -->> HRO: uploadConfirmation()

    HRO ->> SYS: submitRegistration(allFormData)
    SYS -->> HRO: employeeID + successConfirmation()
"""

SSD_UC06 = """
sequenceDiagram
    actor EMP as Employee
    participant SYS as :System

    EMP ->> SYS: navigateToLeaveManagement()
    SYS -->> EMP: displayLeaveForm(currentBalances)

    EMP ->> SYS: selectLeavePeriod(type, startDate, endDate, reason)
    SYS -->> EMP: workingDaysCalculated + balanceValidated()

    EMP ->> SYS: submitLeaveRequest()
    SYS -->> EMP: requestID + statusPending()
    SYS -->> EMP: notificationSentToHR()
"""

SD_UC01 = """
sequenceDiagram
    actor HRO as HR Officer
    participant RC as :RegisterController
    participant ER as :EmployeeRepository
    participant E  as :Employee
    participant NS as :NotificationService

    HRO ->> RC: registerNewEmployee(employeeData)
    Note over RC: GRASP Controller
    RC  ->> ER: validateDuplicate(nationalID, email)
    Note over ER: GRASP Pure Fabrication
    ER  ->> E:  checkDuplicate(nationalID)
    Note over E: GRASP Information Expert
    E  -->> ER: false (no duplicate)
    ER -->> RC: noDuplicate
    RC  ->> ER: create(employeeData)
    ER  ->> E:  save(employee)
    E  -->> ER: employeeID
    ER -->> RC: employeeID
    RC  ->> NS: sendWelcomeEmail(employee)
    Note over NS: GRASP Pure Fabrication
    NS -->> RC: emailQueued
    RC -->> HRO: successConfirmation(employeeID)
"""

SD_UC02 = """
sequenceDiagram
    actor HRO as HR Officer
    participant UC as :UpdateController
    participant ES as :EmployeeService
    participant PUB as :HREventPublisher
    participant ALO as :AuditLogObserver
    participant NO  as :NotificationObserver
    participant PN  as :PayrollNotifier

    HRO  ->> UC:  updateEmployeeRecord(id, updatedData, reason)
    Note over UC: GRASP Controller + GoF Observer Publisher
    UC   ->> ES:  updateEmployee(updatedValues)
    Note over ES: GRASP Information Expert
    ES  -->> UC:  updated
    UC   ->> PUB: publishEvent(EMPLOYEE_UPDATED, id, payload)
    Note over PUB: GoF Observer Subject
    PUB  ->> ALO: onEvent(EMPLOYEE_UPDATED, id, payload)
    Note over ALO: GoF Observer – writes audit log
    PUB  ->> NO:  onEvent(EMPLOYEE_UPDATED, id, payload)
    Note over NO: GoF Observer – sends notification
    UC   ->> PN:  notifyPayroll(id, designation, salary)
    Note over PN: GoF Adapter → ExternalPayrollSystem
    UC  -->> HRO: updateConfirmed
"""

CLASS_DIAGRAM = """
classDiagram
    %% Abstract base (Inheritance + Polymorphism)
    class Report {
        <<abstract>>
        #int id
        #LocalDateTime generatedAt
        +getReportType()* String
        +getSummary()* String
        +getFormat()* String
    }
    class ComplianceReport {
        -String parameters
        -String archivePath
        +getReportType() String
        +getSummary() String
        +getFormat() String
    }
    class HRAnalyticsReport {
        -String reportPeriod
        -String summaryMetrics
        +getReportType() String
        +getSummary() String
        +getFormat() String
    }
    Report <|-- ComplianceReport
    Report <|-- HRAnalyticsReport

    %% GoF Strategy
    class ReportFormatStrategy {
        <<interface>>
        +format(data) String
        +getFormatName() String
    }
    class PdfFormatStrategy { +format() String }
    class ExcelFormatStrategy { +format() String }
    class CsvFormatStrategy { +format() String }
    ReportFormatStrategy <|.. PdfFormatStrategy
    ReportFormatStrategy <|.. ExcelFormatStrategy
    ReportFormatStrategy <|.. CsvFormatStrategy
    class ReportGenerator {
        -Map~String,ReportFormatStrategy~ strategies
        +generateReport() String
        +generateAndFormat() String
    }
    ReportGenerator --> ReportFormatStrategy

    %% GoF Observer
    class HREventObserver {
        <<interface>>
        +onEvent(type, id, payload)
    }
    class HREventPublisher {
        -List~HREventObserver~ observers
        +register(observer)
        +publishEvent(type, id, payload)
    }
    class AuditLogObserver { +onEvent() }
    class NotificationObserver { +onEvent() }
    HREventObserver <|.. AuditLogObserver
    HREventObserver <|.. NotificationObserver
    HREventPublisher --> HREventObserver

    %% GoF Factory Method
    class ReportCreator {
        <<abstract>>
        +createReport(period, userId)* Report
        +generate(period, userId) Report
    }
    class ComplianceReportCreator { +createReport() Report }
    class AnalyticsReportCreator { +createReport() Report }
    ReportCreator <|-- ComplianceReportCreator
    ReportCreator <|-- AnalyticsReportCreator

    %% GoF Adapter
    class PayrollIntegration {
        <<interface>>
        +queueSalaryUpdate()
        +recordHours()
        +initiateSettlement()
    }
    class PayrollSystemAdapter {
        -ExternalPayrollSystem ext
        +queueSalaryUpdate()
        +recordHours()
        +initiateSettlement()
    }
    class ExternalPayrollSystem {
        +submitSalaryChange()
        +logAttendanceHours()
        +processFinalPayout()
    }
    PayrollIntegration <|.. PayrollSystemAdapter
    PayrollSystemAdapter --> ExternalPayrollSystem

    %% GoF Singleton
    class DatabaseConnection {
        -static instance
        -Connection connection
        +getInstance()$ DatabaseConnection
        +getConnection() Connection
    }
"""

PACKAGE_DIAGRAM = """
graph TB
    subgraph PL[" Presentation Layer — com.hr.ui "]
        UI1[EmployeeController]
        UI2[LeaveRequestController]
        UI3[AttendanceController]
        UI4[EvaluationController]
        UI5[ReportController ...]
    end

    subgraph BLL[" Business Logic Layer "]
        subgraph CTRL[" com.hr.controller — SD Controllers "]
            C1[RegisterController UC-01]
            C2[UpdateController UC-02]
            C3[AssignmentController UC-03]
            C4[OffboardingController UC-04]
            C5[LeaveController UC-06]
            C6[ComplianceReportController UC-14]
            C7[HRAnalyticsController UC-15 ...]
        end
        subgraph SVC[" com.hr.service — Services & Patterns "]
            S1[EmployeeService]
            S2[LeaveRequestService]
            S3[PayrollService]
            S4[HREventPublisher GoF Observer]
            S5[ReportGenerator GoF Strategy]
            S6[PayrollSystemAdapter GoF Adapter]
            S7[ReportCreator GoF Factory]
        end
    end

    subgraph DAL[" Data Access Layer — com.hr.dao "]
        D1[EmployeeDAO]
        D2[LeaveRequestDAO]
        D3[DatabaseConnection GoF Singleton]
        D4[Other DAOs ...]
    end

    subgraph MDL[" Domain Model — com.hr.model "]
        M1[Employee]
        M2[Report abstract]
        M3[ComplianceReport]
        M4[HRAnalyticsReport]
        M5[Other Entities ...]
    end

    subgraph DB[" Database "]
        DB1[(MySQL — hr_management 17 tables)]
    end

    PL --> CTRL
    CTRL --> SVC
    SVC --> DAL
    DAL --> DB
    PL --> MDL
    CTRL --> MDL
    SVC --> MDL
    DAL --> MDL
"""

DEPLOYMENT_DIAGRAM = """
graph LR
    subgraph ClientTier["Client Tier"]
        CLIENT[Developer Laptop\nWindows 11\nJava 21 + JavaFX 21]
    end

    subgraph AppTier["Application Tier"]
        APP[Java Application\nMaven Build\nJVM 21]
        subgraph Layers
            L1[Presentation Layer — com.hr.ui]
            L2[Business Logic — com.hr.controller + service]
            L3[Data Access — com.hr.dao]
        end
    end

    subgraph DataTier["Data Tier"]
        DB[(MySQL 8.0.33\nhr_management\n17 tables)]
    end

    CLIENT --> APP
    APP --> DB
    L1 --> L2
    L2 --> L3
    L3 --> DB
"""

COMPONENT_DIAGRAM = """
graph TB
    subgraph UI["UI Components"]
        FXML[14 FXML Layouts]
        CSS[styles.css]
        UICTL[14 JavaFX UI Controllers]
    end

    subgraph BL["Business Logic Components"]
        SDCTL[15 SD Controllers]
        SVCS[16 Service Classes]
        PATTERNS[Design Pattern Classes\nSingleton Observer Strategy\nAdapter Factory Method]
    end

    subgraph DA["Data Access Components"]
        DAOS[14 DAO Classes]
        DBCONN[DatabaseConnection Singleton]
        PROPS[database.properties]
    end

    subgraph EXT["External Systems"]
        MYSQL[(MySQL Server)]
        PAYROLL[External Payroll System\nvia Adapter Pattern]
    end

    FXML --> UICTL
    CSS  --> FXML
    UICTL --> SDCTL
    SDCTL --> SVCS
    SVCS --> PATTERNS
    SVCS --> DAOS
    DAOS --> DBCONN
    DBCONN --> PROPS
    DBCONN --> MYSQL
    PATTERNS --> PAYROLL
"""


# ─────────────────────────────────────────────────────────────
# MAIN DOCUMENT BUILDER
# ─────────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── COVER PAGE ────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Software Requirements and Design Document')
    r.bold = True
    r.font.size = Pt(20)

    doc.add_paragraph()
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run('Smart HR Operations & Workforce Management System')
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()
    for label, value in [
        ('Prepared by:',  'Sohaib Akhlaq (24i-3108) | M. Hasaam (24i-3107) | Shaiman Qasir (24i-3074)'),
        ('Course:',       'SE2002 — Software Design and Architecture'),
        ('Instructor:',   'Ms. Laiba Imran'),
        ('Institution:',  'FAST NUCES, Department of Computer Science'),
        ('Group:',        'Group 16'),
        ('Date:',         datetime.date.today().strftime('%B %d, %Y')),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(label + '  ').bold = True
        p.add_run(value)

    page_break(doc)

    # ── 1. INTRODUCTION ───────────────────────────────────────
    add_heading(doc, '1. Introduction', 1)

    add_heading(doc, '1.1 Purpose', 2)
    add_paragraph(doc,
        'This document specifies the software requirements and design for the Smart HR Operations '
        '& Workforce Management System developed as part of SE2002 Software Design and Architecture '
        'at FAST NUCES (Spring 2026). It covers all 15 use cases across three modules: Employee '
        'Management, Leave & Attendance Management, and Performance & Compliance Management. '
        'The system is implemented in Java 21 with JavaFX 21 for the GUI and MySQL 8.0 for persistence.')

    add_heading(doc, '1.2 Product Scope', 2)
    add_paragraph(doc,
        'The Smart HR Operations & Workforce Management System is a role-based, centralised enterprise '
        'solution that automates HR activities for an organisation. It replaces manual processes and '
        'spreadsheets with structured workflows, policy enforcement, and centralized workforce visibility.')
    add_paragraph(doc, 'Key benefits:', bold=True)
    for b in [
        'Eliminates manual duplication errors in employee records',
        'Enforces leave balance and notice-period policies automatically',
        'Provides real-time attendance tracking and correction workflows',
        'Generates compliance and analytics reports with 5-year archival',
        'Maintains complete audit trail for all HR actions',
    ]:
        add_bullet(doc, b)

    add_heading(doc, '1.3 Objectives', 2)
    for i, obj in enumerate([
        'Automate the end-to-end employee lifecycle from onboarding to offboarding.',
        'Enforce organisational policies (leave quotas, department capacity, probation periods).',
        'Provide role-based access control (Admin / HR Officer / Employee).',
        'Maintain tamper-proof audit logs for all record changes.',
        'Generate compliance and HR analytics reports with long-term archival.',
        'Demonstrate GRASP and GoF design patterns in a real-world system.',
    ], 1):
        add_numbered(doc, obj)

    add_heading(doc, '1.4 Problem Statement', 2)
    add_paragraph(doc,
        'Organisations relying on manual HR processes face delayed approvals, data discrepancies, '
        'policy violations, and lack of transparency. HR officers waste significant time on paperwork, '
        'leave calculations, and attendance reconciliation instead of strategic HR activities.')
    add_paragraph(doc,
        'This system addresses these pain points by providing a centralised JavaFX desktop application '
        'backed by a MySQL database. All business rules (leave balances, department headcounts, notice '
        'periods, audit logging) are enforced automatically by the business logic layer, with no '
        'possibility of bypassing policy through the UI.')

    page_break(doc)

    # ── 2. OVERALL DESCRIPTION ────────────────────────────────
    add_heading(doc, '2. Overall Description', 1)

    add_heading(doc, '2.1 Product Perspective', 2)
    add_paragraph(doc,
        'The system is a standalone desktop application following a 3-tier Layered Architecture:')
    add_bullet(doc, 'Presentation Layer — JavaFX 21 + FXML (14 screens, 14 UI controllers)')
    add_bullet(doc, 'Business Logic Layer — Java SD Controllers (15) + Service classes (16) + Design Pattern classes')
    add_bullet(doc, 'Data Access Layer — 14 DAO classes + MySQL 8.0.33 (17 tables)')
    add_paragraph(doc,
        'The system integrates with an external payroll engine via the Adapter pattern '
        '(PayrollSystemAdapter → ExternalPayrollSystem), shielding HR logic from vendor API changes.')

    add_heading(doc, '2.2 Product Functions', 2)
    for func in [
        'Register, update, and manage employee records with duplicate checking and audit logging',
        'Assign employees to departments with capacity validation',
        'Initiate and manage offboarding workflows with notice-period enforcement',
        'Process leave requests with balance validation and approval/rejection workflow',
        'Record and correct daily attendance with check-in/out time tracking',
        'View real-time dashboard of leave balances and attendance KPIs',
        'Manage performance evaluation cycles with automated reminders',
        'Generate compliance and HR analytics reports with 5-year retention enforcement',
        'Role-based access: Admin configures; HR manages; Employees self-service',
    ]:
        add_bullet(doc, func)

    add_heading(doc, '2.3 List of Use Cases', 2)
    uc_rows = [
        ('UC-01', 'Register New Employee',                'HR Officer',     'Employee Management'),
        ('UC-02', 'Update Employee Record',               'HR Officer',     'Employee Management'),
        ('UC-03', 'Assign Employee to Department',        'HR/Admin',       'Employee Management'),
        ('UC-04', 'Initiate Employee Offboarding',        'HR Officer',     'Employee Management'),
        ('UC-05', 'View Employee Profile',                'Employee/HR',    'Employee Management'),
        ('UC-06', 'Submit Leave Request',                 'Employee',       'Leave & Attendance'),
        ('UC-07', 'Approve or Reject Leave Request',      'HR Officer',     'Leave & Attendance'),
        ('UC-08', 'Record Daily Attendance',              'Employee',       'Leave & Attendance'),
        ('UC-09', 'Request Attendance Correction',        'Employee',       'Leave & Attendance'),
        ('UC-10', 'View Leave Balance & Dashboard',       'Employee/HR',    'Leave & Attendance'),
        ('UC-11', 'Initiate Performance Evaluation Cycle','Admin',          'Performance & Compliance'),
        ('UC-12', 'Submit Employee Performance Evaluation','HR/Evaluator',  'Performance & Compliance'),
        ('UC-13', 'Monitor Employee Probation Status',    'HR Officer',     'Performance & Compliance'),
        ('UC-14', 'Generate Compliance Report',           'Admin',          'Performance & Compliance'),
        ('UC-15', 'Generate HR Analytics Report',         'Admin',          'Performance & Compliance'),
    ]
    add_simple_table(doc, ['UC ID', 'Use Case Name', 'Primary Actor', 'Module'], uc_rows)

    add_heading(doc, '2.4 Extended Use Cases', 2)

    for uc_id, uc_name, actor, pre, post, main_flow, extensions in [
        (
            'UC-01', 'Register New Employee', 'HR Officer',
            'HR authenticated; no duplicate National ID/email; department & designation exist; contract signed.',
            'Unique Employee ID generated (EMP-YYYY-NNNN); record persisted; welcome email sent.',
            [
                'HR navigates to Register New Employee → form displayed with mandatory fields.',
                'HR enters personal details → system checks for duplicate NID and email in real time.',
                'HR selects department, designation, employment type, joining date → probation period auto-populated.',
                'HR uploads documents (PDF/JPEG ≤5 MB) → format and size validated.',
                'HR submits → Employee ID generated, record saved, welcome email dispatched.',
            ],
            ['2a. Duplicate NID → block submission, show error.', '3a. Department at capacity → require HR override.']
        ),
        (
            'UC-06', 'Submit Leave Request', 'Employee',
            'Employee authenticated and active; sufficient leave balance; dates within active policy calendar.',
            'Request stored as Pending Approval; HR notified; dates blocked in attendance calendar.',
            [
                'Employee navigates to Leave Management → form shows current balances by type.',
                'Employee selects leave type, date range, optional reason → working days calculated, balance validated.',
                'Employee submits → status set to Pending Approval, HR notified.',
            ],
            ['2a. Insufficient balance → block with message showing remaining days.', '2b. Blackout period conflict → highlight and suggest adjusted dates.']
        ),
        (
            'UC-07', 'Approve or Reject Leave Request', 'HR Officer',
            'Pending Approval request exists; HR authorised.',
            'Approved → balance deducted + calendar confirmed + employee notified. Rejected → reason recorded + employee notified.',
            [
                'HR opens leave queue → selects pending request → views employee details and leave summary.',
                'HR reviews coverage calendar.',
                'HR approves or rejects with mandatory comment (≥10 characters) → system updates status, deducts balance if approved, notifies employee.',
            ],
            ['3a. Approval would leave department below minimum coverage → warning displayed.']
        ),
    ]:
        add_heading(doc, f'{uc_id} — {uc_name}', 3)
        tbl = add_simple_table(doc,
            ['Field', 'Value'],
            [
                ('Actor', actor),
                ('Preconditions', pre),
                ('Postconditions', post),
            ])
        p = doc.add_paragraph()
        p.add_run('Main Flow:').bold = True
        for i, step in enumerate(main_flow, 1):
            add_numbered(doc, step)
        p2 = doc.add_paragraph()
        p2.add_run('Extensions:').bold = True
        for ext in extensions:
            add_bullet(doc, ext)

    add_heading(doc, '2.5 Use Case Diagram', 2)
    add_diagram_placeholder(doc,
        'Use Case Diagram — All 15 UCs and 3 Actors',
        'UML Use Case Diagram',
        USE_CASE_DIAGRAM)

    page_break(doc)

    # ── 3. NON-FUNCTIONAL REQUIREMENTS ───────────────────────
    add_heading(doc, '3. Non-Functional Requirements', 1)

    add_heading(doc, '3.1 Performance Requirements', 2)
    add_paragraph(doc,
        'NFR-1 (Implemented): Form submissions must complete in < 3 seconds under normal load. '
        'Enforced by PerformanceMonitor service (com.hr.service.PerformanceMonitor) which wraps '
        'timed operations, logs elapsed time, and prints a [PERFORMANCE WARNING] SLA BREACH message '
        'if the 3-second threshold is exceeded.')
    add_bullet(doc, 'Profile page loads < 2 seconds')
    add_bullet(doc, 'Assignment access-rights sync < 5 minutes')

    add_heading(doc, '3.2 Security Requirements', 2)
    add_bullet(doc, 'Role-based access control: Admin / HR Officer / Employee roles enforced via AccessControlService')
    add_bullet(doc, 'All SQL uses PreparedStatement with parameterised queries (SQL injection prevention)')
    add_bullet(doc, 'Passwords stored as hashed values (password_hash column in user_accounts table)')
    add_bullet(doc, 'Read access to sensitive employee fields logged in audit_log table')
    add_bullet(doc, 'Access revocation < 1 hour for involuntary terminations (OffboardingController triggers UserAccount deactivation)')

    add_heading(doc, '3.3 Data Retention Requirements', 2)
    add_paragraph(doc,
        'NFR-2 (Implemented): Compliance reports must be archived for a minimum of 5 years before '
        'deletion is permitted. Enforced by ReportArchive service (com.hr.service.ReportArchive): '
        'archiveReport() registers the 5-year expiry date; canDelete() blocks premature deletion '
        'and logs [POLICY VIOLATION] if attempted within the retention window.')

    add_heading(doc, '3.4 Software Quality Attributes', 2)
    nfr_rows = [
        ('Maintainability', 'High Cohesion + Low Coupling (GRASP) — each class has a single responsibility'),
        ('Extensibility',   'Open/Closed: Strategy pattern allows new report formats without modifying ReportGenerator'),
        ('Testability',     '3-tier separation: business logic in Service/Controller layer is independently testable'),
        ('Reliability',     'Fail-fast validation in all service classes; database constraints enforce referential integrity'),
        ('Auditability',    'AuditLogObserver captures every state change; read access logged via ProfileController'),
    ]
    add_simple_table(doc, ['Attribute', 'How Achieved'], nfr_rows)

    add_heading(doc, '3.5 Business Rules', 2)
    br_rows = [
        ('BR-01', 'Employee ID format: EMP-YYYY-NNNN (generated on registration)'),
        ('BR-02', 'National ID and email must be unique across all employees'),
        ('BR-03', 'Department assignment blocked if current_headcount >= max_headcount'),
        ('BR-04', 'Resignation offboarding requires ≥30 days notice period'),
        ('BR-05', 'Leave approval comment must be ≥10 characters'),
        ('BR-06', 'Annual leave: 21 days | Sick leave: 14 days | Personal: 7 days'),
        ('BR-07', 'Payroll: Tax = 15% of basic; Benefits = 5% of basic; Net = Basic - Tax - Benefits'),
        ('BR-08', 'Performance evaluation score must be 0–100; remarks must not be blank'),
        ('BR-09', 'Compliance reports retained for minimum 5 years (NFR enforcement)'),
    ]
    add_simple_table(doc, ['Rule ID', 'Description'], br_rows)

    add_heading(doc, '3.6 Operating Environment', 2)
    env_rows = [
        ('OS',        'Windows 11 (development); any OS with JVM 21'),
        ('Language',  'Java 21'),
        ('GUI',       'JavaFX 21 + FXML'),
        ('Build',     'Apache Maven 3.x'),
        ('Database',  'MySQL 8.0.33'),
        ('JDBC',      'mysql-connector-java 8.0.33'),
        ('IDE',       'IntelliJ IDEA'),
    ]
    add_simple_table(doc, ['Component', 'Version / Notes'], env_rows)

    page_break(doc)

    # ── 4. USER INTERFACES ───────────────────────────────────
    add_heading(doc, '4. User Interfaces', 1)
    add_paragraph(doc,
        'The system provides 14 JavaFX screens, one per functional module. '
        'Navigation is sidebar-based via MainController.')

    ui_rows = [
        ('employee.fxml',              'EmployeeController',       'UC-01, UC-02, UC-05'),
        ('department.fxml',            'DepartmentController',     'UC-03'),
        ('payroll.fxml',               'PayrollController',        'Payroll Processing'),
        ('leave_request.fxml',         'LeaveRequestController',   'UC-06, UC-07'),
        ('attendance.fxml',            'AttendanceController',     'UC-08'),
        ('attendance_correction.fxml', 'CorrectionController',     'UC-09'),
        ('dashboard.fxml',             'DashboardController',      'UC-10'),
        ('employee_profile.fxml',      'ProfileScreenController',  'UC-05'),
        ('offboarding.fxml',           'OffboardingScreenController','UC-04'),
        ('performance.fxml',           'EvaluationController',     'UC-11, UC-12'),
        ('probation.fxml',             'ProbationController',      'UC-13'),
        ('compliance_report.fxml',     'ReportController',         'UC-14'),
        ('analytics.fxml',             'AnalyticsController',      'UC-15'),
        ('main.fxml',                  'MainController',           'Navigation / Sidebar'),
    ]
    add_simple_table(doc, ['FXML File', 'UI Controller', 'Use Cases'], ui_rows)

    page_break(doc)

    # ── 5. DOMAIN MODEL ──────────────────────────────────────
    add_heading(doc, '5. Domain Model', 1)
    add_paragraph(doc,
        'The domain model contains 14 entity classes in com.hr.model. '
        'ComplianceReport and HRAnalyticsReport inherit from the abstract Report base class, '
        'demonstrating Inheritance and Polymorphism (OOP principles).')
    add_diagram_placeholder(doc,
        'Domain Model — All Entities and Relationships',
        'UML Class Diagram (Domain Model)',
        DOMAIN_MODEL)

    page_break(doc)

    # ── 6. SYSTEM SEQUENCE DIAGRAMS ──────────────────────────
    add_heading(doc, '6. System Sequence Diagrams (SSDs)', 1)
    add_paragraph(doc,
        'SSDs show interactions between actors and the system boundary (black-box view). '
        'Three representative SSDs are shown below.')

    add_heading(doc, '6.1 SSD — UC-01 Register New Employee', 2)
    add_diagram_placeholder(doc, 'SSD UC-01: Register New Employee', 'UML Sequence Diagram (SSD)', SSD_UC01)

    add_heading(doc, '6.2 SSD — UC-06 Submit Leave Request', 2)
    add_diagram_placeholder(doc, 'SSD UC-06: Submit Leave Request', 'UML Sequence Diagram (SSD)', SSD_UC06)

    page_break(doc)

    # ── 7. SEQUENCE DIAGRAMS (INTERNAL) ──────────────────────
    add_heading(doc, '7. Sequence Diagrams (Internal — with GRASP Annotations)', 1)
    add_paragraph(doc,
        'Internal SDs show the class-level interactions within the system. '
        'GRASP pattern roles are annotated on each participant.')

    add_heading(doc, '7.1 SD — UC-01 Register New Employee', 2)
    add_diagram_placeholder(doc, 'SD UC-01: Register New Employee (GRASP annotated)', 'UML Sequence Diagram (Internal)', SD_UC01)

    add_heading(doc, '7.2 SD — UC-02 Update Employee Record (Observer Pattern)', 2)
    add_diagram_placeholder(doc, 'SD UC-02: Update Employee Record with Observer', 'UML Sequence Diagram (Internal)', SD_UC02)

    page_break(doc)

    # ── 8. CLASS DIAGRAM ─────────────────────────────────────
    add_heading(doc, '8. Class Diagram', 1)
    add_paragraph(doc,
        'The class diagram shows all design pattern implementations: '
        'Singleton (DatabaseConnection), Strategy (ReportFormatStrategy), '
        'Observer (HREventPublisher), Factory Method (ReportCreator), '
        'Adapter (PayrollSystemAdapter), and the Inheritance hierarchy (Report → ComplianceReport/HRAnalyticsReport).')
    add_diagram_placeholder(doc,
        'Full Class Diagram — Design Patterns + Inheritance',
        'UML Class Diagram',
        CLASS_DIAGRAM)

    page_break(doc)

    # ── 9. HIGH-LEVEL ARCHITECTURE ───────────────────────────
    add_heading(doc, '9. High-Level Architecture', 1)
    add_paragraph(doc,
        'The system uses a 3-Tier Layered Architecture. This style was chosen because:')
    add_bullet(doc, 'Separation of concerns: UI never accesses the database; business rules never appear in the presentation layer.')
    add_bullet(doc, 'Testability: Service and DAO layers can be tested independently of JavaFX.')
    add_bullet(doc, 'Maintainability: Swapping the database (e.g. MySQL → PostgreSQL) requires changing only the DAO layer.')
    add_bullet(doc, 'Scalability path: The service layer can later be exposed as REST endpoints without changing business logic.')

    add_heading(doc, '9.1 Package Diagram', 2)
    add_diagram_placeholder(doc, 'Package Diagram — 3-Tier Layered Architecture', 'UML Package Diagram', PACKAGE_DIAGRAM)

    add_heading(doc, '9.2 Deployment Diagram', 2)
    add_diagram_placeholder(doc, 'Deployment Diagram', 'UML Deployment Diagram', DEPLOYMENT_DIAGRAM)

    add_heading(doc, '9.3 Component Diagram', 2)
    add_diagram_placeholder(doc, 'Component Diagram', 'UML Component Diagram', COMPONENT_DIAGRAM)

    page_break(doc)

    # ── 10. DESIGN PATTERNS ──────────────────────────────────
    add_heading(doc, '10. Design Patterns', 1)

    add_heading(doc, '10.1 GRASP Patterns', 2)
    grasp_rows = [
        ('Controller',          'RegisterController, UpdateController, etc.',       'Handles system events; delegates to services. Keeps business logic out of UI.'),
        ('Information Expert',  'EmployeeService, LeaveRequestService, PayrollService', 'Class with data owns the responsibility. PayrollService owns tax calculation.'),
        ('Creator',             'LeaveRequestService, OffboardingController',       'LeaveRequestService creates LeaveRequest objects; owns all construction data.'),
        ('Low Coupling',        '3-tier architecture',                              'UI depends only on SD Controllers; Controllers only on Services; Services only on DAOs.'),
        ('High Cohesion',       'AuditLogService, NotificationService',            'Each service handles exactly one concern. AuditLogService only logs; NS only notifies.'),
        ('Pure Fabrication',    'EmployeeRepository, NotificationService, AuditLogService, PayrollNotifier', 'Artificial classes for cross-cutting concerns that do not represent domain entities.'),
        ('Protected Variation', 'DatabaseConnection (Singleton)',                   'Shields all 14 DAOs from JDBC connection complexity. Change only DatabaseConnection if JDBC setup changes.'),
        ('Indirection',         'Service layer between UI and DAO',                 'UI never calls DAO directly; service layer acts as intermediary, enabling independent evolution.'),
    ]
    add_simple_table(doc, ['Pattern', 'Class(es)', 'Justification'], grasp_rows)

    add_heading(doc, '10.2 GoF Patterns', 2)
    gof_rows = [
        ('Singleton',       'DatabaseConnection',                                       'Ensures one shared JDBC connection instance. Prevents connection pool exhaustion.'),
        ('Strategy',        'ReportFormatStrategy → PdfFormatStrategy, ExcelFormatStrategy, CsvFormatStrategy', 'Format selection at runtime without if-else in ReportGenerator. Adding new formats requires only a new class (OCP).'),
        ('Observer',        'HREventPublisher → AuditLogObserver, NotificationObserver', 'UC-02 publishes EMPLOYEE_UPDATED; observers react independently. New reactions (analytics, payroll) add without changing UpdateController.'),
        ('Factory Method',  'abstract ReportCreator → ComplianceReportCreator, AnalyticsReportCreator', 'Subclasses decide which Report subtype to instantiate. Caller works with abstract ReportCreator type.'),
        ('Adapter',         'PayrollSystemAdapter wraps ExternalPayrollSystem, implements PayrollIntegration', 'HR domain calls (int employeeId) translated to legacy API calls (String empRef). Vendor change = one class change.'),
    ]
    add_simple_table(doc, ['Pattern', 'Class(es)', 'Justification'], gof_rows)

    page_break(doc)

    # ── 11. OOP PRINCIPLES ───────────────────────────────────
    add_heading(doc, '11. OOP Principles', 1)
    oop_rows = [
        ('Encapsulation',  'All 14 model classes have private fields with public getters/setters. DAO layer hides SQL. Services hide DAO details.',
         'Employee.java — all fields private\nEmployeeService.java — wraps EmployeeDAO'),
        ('Abstraction',    'Controllers expose single-method workflows hiding 5+ internal steps. Services abstract business rules from UI.',
         'RegisterController.registerNewEmployee() — 1 call, 5 internal steps'),
        ('Inheritance',    'abstract Report ← ComplianceReport, HRAnalyticsReport. abstract ReportCreator ← ComplianceReportCreator, AnalyticsReportCreator.',
         'model/Report.java\nservice/ReportCreator.java'),
        ('Polymorphism',   'Report.getReportType(), getSummary(), getFormat() overridden in each subclass. ReportCreator.createReport() overridden in each factory. ReportFormatStrategy.format() overridden in each strategy.',
         'All abstract method overrides in ComplianceReport, HRAnalyticsReport'),
    ]
    add_simple_table(doc, ['Principle', 'How Implemented', 'Key Files'], oop_rows)

    page_break(doc)

    # ── 12. DATABASE ─────────────────────────────────────────
    add_heading(doc, '12. Database Design', 1)
    add_paragraph(doc,
        'The system uses MySQL 8.0.33. The schema (database/schema.sql) contains 17 normalised tables '
        'with foreign key constraints, UNIQUE constraints, and ENUM columns for status fields.')

    db_rows = [
        ('departments',            '7',  'id (PK), name (UNIQUE), manager_id, parent_dept_id, max_headcount, current_headcount'),
        ('designations',           '4',  'id (PK), title, description, salary_grade'),
        ('employees',              '15', 'id (PK), email (UNIQUE), national_id (UNIQUE), department_id (FK), designation_id (FK), status ENUM'),
        ('user_accounts',          '7',  'employee_id (UNIQUE FK), username (UNIQUE), password_hash, role ENUM(ADMIN/HR/EMPLOYEE)'),
        ('payroll',                '9',  'employee_id (FK), pay_period_start, pay_period_end, basic_salary, tax_deduction, net_pay'),
        ('leave_requests',         '11', 'employee_id (FK), leave_type ENUM, start_date, end_date, status ENUM(PENDING/APPROVED/REJECTED)'),
        ('leave_balances',         '5',  'employee_id (UNIQUE FK), annual_balance (21), sick_balance (14), personal_balance (7)'),
        ('attendance_records',     '8',  'employee_id (FK), attendance_date, check_in_time, check_out_time, total_hours, attendance_status ENUM'),
        ('attendance_corrections', '9',  'attendance_id (FK), employee_id (FK), original_value, corrected_value, status ENUM'),
        ('offboarding_workflows',  '9',  'employee_id (FK), separation_type ENUM, last_working_date, status ENUM'),
        ('probation_records',      '8',  'employee_id (FK), start_date, end_date, extensions, decision ENUM'),
        ('evaluation_cycles',      '7',  'cycle_name, start_date, end_date, evaluation_type, status ENUM'),
        ('performance_evaluations','8',  'employee_id (FK), evaluator_id (FK), cycle_id (FK), aggregate_score, status ENUM'),
        ('compliance_reports',     '8',  'report_type, generated_at, format ENUM(PDF/Excel), status, archive_path'),
        ('hr_analytics_reports',   '5',  'creator_id (FK), report_period, summary_metrics'),
        ('audit_log',              '9',  'user_id, action, entity_type, entity_id, field_name, old_value, new_value, timestamp'),
        ('org_history',            '5',  'employee_id (FK), old_dept_id (FK), new_dept_id (FK), effective_date'),
    ]
    add_simple_table(doc, ['Table Name', 'Columns', 'Key Columns'], db_rows)

    add_paragraph(doc,
        '\nData Flow: UI Controller → SD Controller → Service → DAO → MySQL\n'
        'Example (UC-01): EmployeeController.handleAdd() → RegisterController.registerNewEmployee() '
        '→ EmployeeRepository.validateDuplicate() → EmployeeDAO.existsByNationalId() → SELECT on employees table.')

    page_break(doc)

    # ── 13. WORK DIVISION ─────────────────────────────────────
    add_heading(doc, '13. Work Division Table', 1)
    work_rows = [
        ('Sohaib Akhlaq\n24i-3108',
         'UC-01 Register Employee, UC-02 Update Record, UC-03 Assign to Department, UC-04 Offboarding, UC-05 View Profile',
         'RegisterController, UpdateController, AssignmentController, OffboardingController, ProfileController\nEmployeeDAO, DepartmentDAO, OffboardingWorkflowDAO\nemployee.fxml, department.fxml, offboarding.fxml, employee_profile.fxml\nGoF Observer pattern, Adapter pattern'),
        ('M. Hasaam\n24i-3107',
         'UC-11 Evaluation Cycle, UC-12 Performance Eval, UC-13 Probation, UC-14 Compliance Report, UC-15 HR Analytics',
         'EvaluationCycleController, PerformanceEvalController, ProbationMonitorController, ComplianceReportController, HRAnalyticsController\nEvaluationCycleDAO, PerformanceEvaluationDAO, ProbationRecordDAO, ComplianceReportDAO, HRAnalyticsReportDAO\nperformance.fxml, probation.fxml, compliance_report.fxml, analytics.fxml\nGoF Factory Method, Strategy, abstract Report hierarchy'),
        ('Shaiman Qasir\n24i-3074',
         'UC-06 Submit Leave, UC-07 Approve Leave, UC-08 Record Attendance, UC-09 Attendance Correction, UC-10 Dashboard',
         'LeaveController, LeaveApprovalController, AttendanceRecordController, AttendanceCorrectionController, LeaveDashboardController\nLeaveRequestDAO, AttendanceRecordDAO, AttendanceCorrectionDAO\nleave_request.fxml, attendance.fxml, attendance_correction.fxml, dashboard.fxml\nDatabase schema, PayrollService'),
    ]
    add_simple_table(doc, ['Member', 'Use Cases', 'Code Deliverables'], work_rows)

    doc.add_paragraph()
    add_paragraph(doc,
        'Note: All code is integrated into a single Maven project (one pom.xml, one codebase). '
        'The above division reflects the original assignment allocation; all 15 UCs are fully implemented '
        'and tested within the unified system.',
        italic=True)

    # ── SAVE ─────────────────────────────────────────────────
    output_path = 'docs/project_report.docx'
    doc.save(output_path)
    print('\n[DONE] Report saved to: ' + output_path)
    print('    Sections: Introduction, Overall Description, NFRs, UI, Domain Model,')
    print('              SSDs, SDs, Class Diagram, Architecture (Package/Deployment/Component),')
    print('              Design Patterns, OOP Principles, Database, Work Division')
    print('\n[DIAGRAMS] 9 Mermaid placeholders included.')
    print('    To add diagram images: open project_report.docx in Claude.ai and ask:')
    print('    "Replace every DIAGRAM PLACEHOLDER block with the rendered diagram image."')


if __name__ == '__main__':
    build_document()
