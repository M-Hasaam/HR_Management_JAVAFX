"""
Converts Project_Report_Final.md to a formatted Word .docx document.
Run: python generate_report.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

MD_PATH  = Path(__file__).parent / "Project_Report_Final.md"
OUT_PATH = Path(__file__).parent / "Project_Report_Final.docx"

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x1A, 0x37, 0x5E)   # headings
TEAL        = RGBColor(0x00, 0x7B, 0x83)   # h2
STEEL       = RGBColor(0x2E, 0x4D, 0x6B)   # h3
CODE_BG     = RGBColor(0xF4, 0xF4, 0xF4)
CODE_BORDER = RGBColor(0xCC, 0xCC, 0xCC)
TABLE_HDR   = RGBColor(0x1A, 0x37, 0x5E)
TABLE_ROW   = RGBColor(0xEA, 0xF1, 0xFB)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT   = RGBColor(0x1C, 0x1C, 0x1C)
MID_GRAY    = RGBColor(0x55, 0x55, 0x55)


def rgb_hex(rgb: RGBColor) -> str:
    """Return the 6-char uppercase hex string from an RGBColor."""
    return str(rgb).upper()  # RGBColor.__str__ already returns e.g. '1A375E'


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  rgb_hex(rgb))
    tcPr.append(shd)


def add_border(paragraph, rgb: RGBColor):
    """Add a left border to a paragraph (used for code blocks)."""
    pPr   = paragraph._p.get_or_add_pPr()
    pBdr  = OxmlElement('w:pBdr')
    left  = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    '12')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), rgb_hex(rgb))
    pBdr.append(left)
    pPr.append(pBdr)


def set_para_shading(paragraph, rgb: RGBColor):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  rgb_hex(rgb))
    pPr.append(shd)


def apply_inline_markup(run_parent, text: str, base_font_size=11, bold=False,
                        color: RGBColor = None):
    """
    Parse inline markdown (**bold**, *italic*, `code`) and add runs to
    run_parent (a paragraph).  Returns nothing — modifies in place.
    """
    segments = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for seg in segments:
        if not seg:
            continue
        run = run_parent.add_run()
        if seg.startswith('**') and seg.endswith('**'):
            run.text = seg[2:-2]
            run.bold = True
            run.font.size = Pt(base_font_size)
        elif seg.startswith('*') and seg.endswith('*'):
            run.text = seg[1:-1]
            run.italic = True
            run.font.size = Pt(base_font_size)
        elif seg.startswith('`') and seg.endswith('`'):
            run.text = seg[1:-1]
            run.font.name = 'Courier New'
            run.font.size = Pt(base_font_size - 1)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run.text = seg
            run.font.size = Pt(base_font_size)
        if bold and not seg.startswith('**'):
            run.bold = bold
        if color:
            run.font.color.rgb = color


def add_title_page(doc: Document, lines: list[str]):
    """Build a styled title page from the first few header lines."""
    doc.add_picture  # placeholder — we do it manually below

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(72)
    t.paragraph_format.space_after  = Pt(6)
    run = t.add_run("Smart HR Operations &\nWorkforce Management System")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # Subtitle
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_after = Pt(4)
    r2 = s.add_run("Final Project Report — SE2002 Software Design and Architecture")
    r2.font.size = Pt(14)
    r2.font.color.rgb = TEAL

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.paragraph_format.space_after = Pt(36)
    r3 = sub2.add_run("Spring 2026  |  FAST NUCES, Islamabad  |  Department of Computer Science")
    r3.font.size = Pt(11)
    r3.font.color.rgb = MID_GRAY

    # Info table
    meta = [
        ("Course",          "SE2002 — Software Design and Architecture"),
        ("Instructor",      "Ms. Laiba Imran"),
        ("Group Number",    "Group 16"),
        ("Submission Type", "Solo Implementation (all 15 Use Cases, all 4 Modules)"),
        ("Date",            "April 2026"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    for i, (label, value) in enumerate(meta):
        row = tbl.rows[i]
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.0)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)
        if i % 2 == 0:
            set_cell_bg(row.cells[0], TABLE_ROW)
            set_cell_bg(row.cells[1], TABLE_ROW)

    doc.add_page_break()


def add_toc_placeholder(doc: Document):
    h = doc.add_paragraph("Table of Contents", style='Heading 1')
    p = doc.add_paragraph()
    run = p.add_run("[Right-click in Word → Update Field to refresh the TOC]")
    run.italic = True
    run.font.color.rgb = MID_GRAY
    doc.add_page_break()


def style_heading(paragraph, level: int):
    if level == 1:
        paragraph.runs[0].font.color.rgb = NAVY
        paragraph.runs[0].font.size = Pt(18)
        paragraph.paragraph_format.space_before = Pt(20)
        paragraph.paragraph_format.space_after  = Pt(8)
    elif level == 2:
        paragraph.runs[0].font.color.rgb = TEAL
        paragraph.runs[0].font.size = Pt(14)
        paragraph.paragraph_format.space_before = Pt(16)
        paragraph.paragraph_format.space_after  = Pt(6)
    elif level == 3:
        paragraph.runs[0].font.color.rgb = STEEL
        paragraph.runs[0].font.size = Pt(12)
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after  = Pt(4)
    elif level == 4:
        paragraph.runs[0].font.color.rgb = DARK_TEXT
        paragraph.runs[0].font.size = Pt(11)
        paragraph.runs[0].bold = True
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after  = Pt(3)


def build_md_table(doc: Document, header_row: list[str], rows: list[list[str]]):
    col_count = len(header_row)
    tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr_cells = tbl.rows[0].cells
    for ci, text in enumerate(header_row):
        set_cell_bg(hdr_cells[ci], TABLE_HDR)
        p = hdr_cells[ci].paragraphs[0]
        run = p.add_run(text.strip())
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)

    # Data rows
    for ri, row in enumerate(rows):
        row_cells = tbl.rows[ri + 1].cells
        bg = TABLE_ROW if ri % 2 == 0 else WHITE
        for ci, cell_text in enumerate(row):
            if ci < len(row_cells):
                set_cell_bg(row_cells[ci], bg)
                p = row_cells[ci].paragraphs[0]
                apply_inline_markup(p, cell_text.strip(), base_font_size=10)

    doc.add_paragraph()   # spacing after table


def emit_code_block(doc: Document, code_lines: list[str]):
    for line in code_lines:
        p = doc.add_paragraph()
        set_para_shading(p, CODE_BG)
        add_border(p, CODE_BORDER)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.2)
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = DARK_TEXT
    # small gap after block
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(0)
    gap.paragraph_format.space_after  = Pt(4)


def emit_body_para(doc: Document, text: str, bold=False):
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    apply_inline_markup(p, text, base_font_size=11, bold=bold)


def emit_bullet(doc: Document, text: str, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    apply_inline_markup(p, text.strip(), base_font_size=11)


def emit_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    apply_inline_markup(p, text.strip(), base_font_size=11)


# ── Main parser ───────────────────────────────────────────────────────────────

def convert(md_path: Path, out_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Default body font
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = DARK_TEXT

    raw = md_path.read_text(encoding='utf-8')
    lines = raw.splitlines()

    # Title page — extract from top of file
    add_title_page(doc, lines)
    add_toc_placeholder(doc)

    # State machine
    in_code      = False
    code_lines   = []
    in_table     = False
    tbl_header   = []
    tbl_rows     = []
    skip_until   = 0   # line index to resume at after skipping title block

    # Find end of title block (first "---" separator)
    for i, ln in enumerate(lines):
        if i > 0 and ln.strip() == '---':
            skip_until = i + 1
            break

    # Skip second metadata block (Course/Instructor lines)
    # We'll detect it heuristically: next "---" after skip_until
    second_end = skip_until
    for i in range(skip_until, len(lines)):
        if lines[i].strip() == '---':
            second_end = i + 1
            break

    # Skip TOC section (## Table of Contents ... blank line before ## 1.)
    in_toc = False
    toc_end = second_end
    for i in range(second_end, len(lines)):
        ln = lines[i].strip()
        if ln.startswith('## Table of Contents'):
            in_toc = True
        if in_toc and ln.startswith('---'):
            toc_end = i + 1
            break

    parse_from = toc_end

    i = parse_from
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Code fence ──────────────────────────────────────────────────────
        if stripped.startswith('```'):
            if not in_code:
                in_code    = True
                code_lines = []
            else:
                emit_code_block(doc, code_lines)
                in_code    = False
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line.rstrip())
            i += 1
            continue

        # ── Horizontal rule ─────────────────────────────────────────────────
        if stripped in ('---', '***', '___') and len(stripped) >= 3:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            run = p.add_run('─' * 72)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(9)
            i += 1
            continue

        # ── Table detection ─────────────────────────────────────────────────
        if stripped.startswith('|') and stripped.endswith('|'):
            # Check if next line is separator
            if not in_table:
                # this is header row
                tbl_header = [c.strip() for c in stripped.split('|')[1:-1]]
                in_table   = True
                tbl_rows   = []
                i += 1
                # skip separator line
                if i < len(lines) and re.match(r'^\|[\s\-|:]+\|$', lines[i].strip()):
                    i += 1
                continue
            else:
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                tbl_rows.append(cells)
                i += 1
                continue
        else:
            if in_table:
                build_md_table(doc, tbl_header, tbl_rows)
                in_table   = False
                tbl_header = []
                tbl_rows   = []

        # ── Headings ─────────────────────────────────────────────────────────
        hm = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if hm:
            level = len(hm.group(1))
            text  = hm.group(2).strip()
            # Remove any markdown link syntax from heading
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

            style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
            p = doc.add_paragraph(style=style_map.get(level, 'Heading 4'))
            run = p.add_run(text)
            style_heading(p, level)
            i += 1
            continue

        # ── Bullet lists ──────────────────────────────────────────────────────
        bm = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if bm:
            indent = len(bm.group(1)) // 2
            emit_bullet(doc, bm.group(2), level=indent)
            i += 1
            continue

        # ── Numbered lists ────────────────────────────────────────────────────
        nm = re.match(r'^\s*\d+\.\s+(.*)', line)
        if nm:
            emit_numbered(doc, nm.group(1))
            i += 1
            continue

        # ── Bold-only line (field label style) ───────────────────────────────
        if stripped.startswith('**') and stripped.endswith('**') and stripped.count('**') == 2:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(stripped[2:-2])
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = NAVY
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── Regular paragraph ─────────────────────────────────────────────────
        emit_body_para(doc, stripped)
        i += 1

    # Flush any unclosed table
    if in_table:
        build_md_table(doc, tbl_header, tbl_rows)

    # ── Footer on every page ─────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(
            "Smart HR Operations & Workforce Management System  |  "
            "SE2002 Final Report  |  FAST NUCES Islamabad  |  Spring 2026"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = MID_GRAY

    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    convert(MD_PATH, OUT_PATH)
